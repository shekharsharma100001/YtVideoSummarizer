import streamlit as st
import google.generativeai as genai 
from youtube_transcript_api import YouTubeTranscriptApi
import os
import re
import time
import io
from fpdf import FPDF
from docx import Document

st.set_page_config(layout="wide")
key = st.secrets["GOOGLE_API_KEY"]
genai.configure(api_key=key)

# Extracts video id from youtube URL
def extract_video_id(youtube_url):
    regex = (
        r'(?:https?://)?(?:www\.)?(?:youtube\.com/watch\?v=|youtu\.be/|youtube\.com/embed/|youtube\.com/v/|youtube\.com/shorts/|youtube.com/playlist\?list=)([^&=%\?]{11})'
    )
    match = re.match(regex, youtube_url)
    if match:
        return match.group(1)
    else:
        raise ValueError("Invalid YouTube URL")

# This function will Extract Transcript from the yt Video
def get_english_transcript(video_id):
    try:
        transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)

        try:
            transcript = transcript_list.find_manually_created_transcript(['en'])
            transcript_text = transcript.fetch()
            transcript = ' '.join([i['text'] for i in transcript_text])
            return transcript

        except:
            pass

        try:
            transcript = transcript_list.find_generated_transcript(['en'])
            transcript_text = transcript.fetch()
            transcript = ' '.join([i['text'] for i in transcript_text])
            return transcript

        except:
            pass

        for transcript in transcript_list:
            if transcript.language_code != 'en' and transcript.is_translatable:
                translated_transcript = transcript.translate('en')
                transcript_text = translated_transcript.fetch()
                transcript = ' '.join([i['text'] for i in transcript_text])
                return transcript

        return None

    except:
        return None

def generate_gemini_content(transcript_text, prompt):
    model = genai.GenerativeModel("gemini-pro")
    response = model.generate_content(prompt + transcript_text)
    return response.text

### Streamlit app
st.image(r"images/yt.png", use_column_width=True)

text = '''
<h1 style="text-align: right;">
    <span style="color:cyan">Video</span> 
    <span style="color:red">Summarizer</span>
</h1>
'''
st.markdown(text, unsafe_allow_html=True)

st.title("📢 :green[Instructions]")
st.info("""
1. Paste the URL of Youtube Video.
2. You can use your own prompt to generate your summary according to requirements.
3. Select the word limits of your Summary by moving the slider.
""")
st.header("🔗 :green[Paste Youtube URL]")
youtube_link = st.text_input('yt link', help="Copy Your Youtube URL and paste it in the Text Box", label_visibility="hidden")
thumbnail, desc = st.columns(2)

if youtube_link:
    video_id = extract_video_id(youtube_link)
    with thumbnail.container(border=True):
        st.image(f"http://img.youtube.com/vi/{video_id}/0.jpg", use_column_width=True)
    with desc.container(border=True):
        st.header(":violet[Word Limit] ⏩")
        with st.container(border = True):
            words = st.slider('word limit', 50, 500, 250, 10, label_visibility='hidden')
        with st.container(border = True):
            prompt = f"""You are a YouTube video summarizer. You will be taking the transcript text 
        and summarizing the entire video, providing the important summary in points within {words} words.
        Please provide the summary of the given YouTube caption here: """
            st.header(":violet[Enter Prompt(Optional)]")
            prompt = st.text_input("prompt", label_visibility='hidden', help='Enter the prompt to clarify in which format summary is to be generated.')
        st.write(' ')
        st.markdown('<h7><h7>',unsafe_allow_html=True)
        emoji, btn, spin = st.columns([4, 3, 3])
        emoji.info("Click Here ➡️")

        if btn.button("Generate Summary", use_container_width=True):
            with spin:
                with st.spinner('Wait...'):
                    time.sleep(1)
                    st.success('Generating.....')
                transcript_text = get_english_transcript(video_id)
                if transcript_text:
                    summary = generate_gemini_content(transcript_text, prompt)
                    st.session_state['summary'] = summary
                else:
                    st.session_state['summary'] = False

if 'summary' in st.session_state and st.session_state['summary'] == False:
    st.info("🥲 Sorry, Summary for the video cannot be generated. Try with another video.")

if 'summary' in st.session_state and st.session_state['summary']:
    summary = st.session_state['summary']
    with st.container(border=True):
        st.markdown('## Detailed Notes:')
        st.write(summary)

    class PDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 12)
            self.cell(0, 10, 'Summary Report', 0, 1, 'C')

        def footer(self):
            self.set_y(-15)
            self.set_font('Arial', 'I', 8)
            self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

        def add_summary(self, summary):
            self.add_page()
            self.set_font('Arial', '', 12)
            self.multi_cell(0, 10, summary, align='L')

    def load_pdf():
        pdf = PDF()
        pdf.add_summary(summary)
        pdf_file = io.BytesIO(pdf.output(dest='S').encode('latin1'))
        return pdf_file

    def load_docx():
        doc = Document()
        doc.add_heading('Summary Report', 0)
        doc.add_paragraph(summary)
        word_file = io.BytesIO()
        doc.save(word_file)
        word_file.seek(0)
        return word_file

    col1, col2 = st.columns(2)

    with col1:
        with col1.container(border = True):
            st.header(':green[Import as PDF file] ⬇️')
            st.download_button(
                label="Download PDF ✅",
                data=load_pdf(),
                file_name="summary_report.pdf",
                mime="application/pdf"
            )

    with col2:
        with col2.container(border = True):
            st.header(':green[Import as Word file] ⬇️')
            st.download_button(
                label="Download Word ✅",
                data=load_docx(),
                file_name="summary_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
